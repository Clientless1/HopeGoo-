import os,sys,json,re,base64
from datetime import datetime
sys.path.insert(0,'.')
import hopegoo_api as api
import requests,urllib3;urllib3.disable_warnings()

OUT="/Users/congwenjie/Desktop/未命名文件夹";os.makedirs(OUT,exist_ok=True)
DB=os.path.expanduser("~/Library/Application Support/com.reqable.macosx/box/data.mdb")
PROGRESS=os.path.join(OUT,"progress.json")

# 恢复进度
if os.path.exists(PROGRESS):
    p=json.load(open(PROGRESS));all_res=p["results"];done=set(p["done"]);start_idx=p["idx"]
    print(f"📂 恢复: {len(all_res)}家, {len(done)}城, 从第{start_idx+1}城续")
else:all_res=[];done=set();start_idx=0

# 城市列表
text="""北京 天津 石家庄 唐山 秦皇岛 邯郸 邢台 保定 张家口 承德 沧州 廊坊 衡水
太原 大同 阳泉 长治 晋城 朔州 忻州 吕梁 晋中 临汾 运城
呼和浩特 包头 乌海 赤峰 通辽 鄂尔多斯 呼伦贝尔 巴彦淖尔 乌兰察布
沈阳 大连 鞍山 抚顺 本溪 丹东 锦州 营口 阜新 辽阳 盘锦 铁岭 朝阳 葫芦岛
长春 吉林 四平 辽源 通化 白山 松原 白城
哈尔滨 齐齐哈尔 鸡西 鹤岗 双鸭山 大庆 伊春 佳木斯 七台河 牡丹江 黑河 绥化
上海 南京 无锡 徐州 常州 苏州 南通 连云港 淮安 盐城 扬州 镇江 泰州 宿迁
杭州 宁波 温州 嘉兴 湖州 绍兴 金华 衢州 舟山 台州 丽水
合肥 芜湖 蚌埠 淮南 马鞍山 淮北 铜陵 安庆 黄山 滁州 阜阳 宿州 六安 亳州 池州 宣城
福州 厦门 莆田 三明 泉州 漳州 南平 龙岩 宁德
南昌 景德镇 萍乡 九江 新余 鹰潭 赣州 吉安 宜春 抚州 上饶
济南 青岛 淄博 枣庄 东营 烟台 潍坊 济宁 泰安 威海 日照 临沂 德州 聊城 滨州 菏泽
郑州 开封 洛阳 平顶山 安阳 鹤壁 新乡 焦作 濮阳 许昌 漯河 三门峡 南阳 商丘 信阳 周口 驻马店
武汉 黄石 十堰 宜昌 襄阳 鄂州 荆门 孝感 荆州 黄冈 咸宁 随州 仙桃 潜江 天门
长沙 株洲 湘潭 衡阳 邵阳 岳阳 常德 张家界 益阳 郴州 永州 怀化 娄底
广州 深圳 珠海 汕头 佛山 韶关 湛江 茂名 肇庆 惠州 梅州 汕尾 河源 阳江 清远 东莞 中山 潮州 揭阳 云浮
南宁 柳州 桂林 梧州 北海 防城港 钦州 贵港 玉林 百色 贺州 河池 来宾 崇左
海口 三亚 三沙 儋州
重庆 成都 自贡 攀枝花 泸州 德阳 绵阳 广元 遂宁 内江 乐山 南充 眉山 宜宾 广安 达州 雅安 巴中 资阳
贵阳 六盘水 遵义 安顺 毕节 铜仁
昆明 曲靖 玉溪 保山 昭通 丽江 普洱 临沧
拉萨 日喀则 昌都 林芝 那曲
西安 铜川 宝鸡 咸阳 渭南 延安 汉中 榆林 安康 商洛
兰州 嘉峪关 金昌 白银 天水 武威 张掖 平凉 酒泉 庆阳 定西 陇南
西宁 海东 银川 石嘴山 吴忠 固原 中卫
乌鲁木齐 克拉玛依 吐鲁番 哈密 石河子 阿拉尔 五家渠""".split()

CITIES=[];seen=set()
for t in text:
    t=t.strip()
    if t and t not in seen:CITIES.append(t);seen.add(t)
cmap=api.load_city_map()
CITY_LIST=[(c,cmap.get(c) or cmap.get(c.replace("市",""))) for c in CITIES if cmap.get(c) or cmap.get(c.replace("市",""))]

def save():json.dump({"results":all_res,"done":list(done),"idx":i},open(PROGRESS,"w"),ensure_ascii=False)

def fresh_hd():
    raw=open(DB,"rb").read();best=None
    for blk in re.findall(rb'H4sI[A-Za-z0-9+/=]{80,}',raw):
        try:d=json.loads(__import__('gzip').decompress(base64.b64decode(blk+b'=='*(-len(blk)%4))).decode("utf-8","replace"))
        except:continue
        if not isinstance(d,dict):continue
        s=d.get("session",{})
        if s.get("connection",{}).get("originHost","")!="hotel.reshg.com":continue
        p=s.get("request",{}).get("requestLine",{}).get("path","")
        if "/tapi/v2/list?" not in p:continue
        hs=s.get("request",{}).get("headers",[]);low=[str(h).lower() for h in hs]
        if any(l.startswith("sectoken:") for l in low) and any(l.startswith("dun-token:") for l in low):
            ckc=sum(1 for l in low if l.startswith("cookie:"))
            ts=s.get("timestamp",0)
            # 优先App请求(cookie多)，其次最新
            score=(1 if ckc>10 else 0,ts)
            if best is None or score>best[0]:best=(score,hs)
    if not best:raise RuntimeError("无有效token")
    (_,_),path,hs=best;hd={};cookies=[]
    for h in hs:
        s2=str(h);i=s2.find(":")
        if i<0:continue
        nm=s2[:i].strip();v=s2[i+1:].strip()
        if nm.lower()=="cookie":cookies.append(v)
        else:hd[nm]=v
    hd["Cookie"]="; ".join(cookies);return hd,path

s=requests.Session();s.verify=False
hd=None;i=start_idx-1 if start_idx>0 else -1

while True:
    i+=1
    if i>=len(CITY_LIST):
        print("🎉 全部完成!");save();break
    name,cid=CITY_LIST[i]
    if name in done:continue

    if hd is None:
        try:hd,orig_path=fresh_hd()
        except:
            print(f"\n🔴 [{i+1}/{len(CITY_LIST)}] 需App搜一次 → 然后重新运行此脚本")
            save();sys.exit(1)

    try:
        # 用原始请求做模板，只替换 city 和分页
        import urllib.parse as _up
        _parsed=_up.urlparse("https://hotel.reshg.com"+orig_path)
        _params=dict(_up.parse_qsl(_parsed.query))
        _params["city"]=str(cid)
        _params["pageIndex"]="0"
        _params["pageSize"]="20"
        r=s.get("https://hotel.reshg.com/tapi/v2/list?"+_up.urlencode(_params),headers=hd,timeout=15)
        d=r.json();ec=str(d.get("errorCode",""))
        if ec=="-99":
            print(f"\n[{i+1}/{len(CITY_LIST)}] {name} 🔴token过期")
            print("  请App搜一次，然后重新运行: .venv/bin/python3 batch_all.py")
            save();sys.exit(1)
        hl=(d.get("data")or{}).get("hotelList")or[]
        for hh in hl:
            gcoin=next((lb.get("amount") for lb in (hh.get("productLabelList")or[]) if lb.get("productLabelId")==345),None)
            all_res.append({"name":(hh.get("hotelName")or"").strip(),"address":(hh.get("hotelAddress")or"").strip()or"-",
                "couponPrice":hh.get("couponPrice",""),"discountPrice":hh.get("discountPrice",""),"price":hh.get("price",""),
                "gcoinAmount":gcoin,"_city":name})
        done.add(name)
        print(f"[{i+1}/{len(CITY_LIST)}] {name}: {len(hl)}家 (累计{len(all_res)}/{len(done)}城)")
    except Exception as e:
        print(f"[{i+1}] {name}: {e}");hd=None

    if (i+1)%5==0:save()

# 出文档
all_res.sort(key=lambda r:(-(r.get("couponPrice")or 0),(r.get("discountPrice")or 1e9)))
from docx import Document;from docx.shared import Pt;from docx.enum.text import WD_ALIGN_PARAGRAPH;from docx.oxml.ns import qn
def frun(p,b=False,s=9):
    r=p.add_run("");r.bold=b;r.font.size=Pt(s);r.font.name='宋体';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');return r
doc=Document()
st=doc.styles['Normal'];st.font.name='宋体';st.font.size=Pt(10);st.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
t=doc.add_paragraph();t.alignment=WD_ALIGN_PARAGRAPH.CENTER;frun(t,True,18).text="HopeGoo G-Coins 全国酒店"
info=doc.add_paragraph()
for l in [f"生成：{datetime.now():%Y-%m-%d %H:%M}",f"覆盖：{len(done)}城",f"酒店：{len(all_res)}家","排序：抵扣高→低，付款低→高"]:
    frun(info).text=l+"\n"
tb=doc.add_table(rows=1,cols=6);tb.style="Light Grid Accent 1"
for i,c in enumerate(["序号","酒店名称","酒店地址","抵扣金额(CNY)","付款金额(CNY)","房价"]):
    frun(tb.rows[0].cells[i].paragraphs[0],True,10).text=c
for i,x in enumerate(all_res,1):
    cells=tb.add_row().cells
    gcoin=x.get("gcoinAmount")  # G-Coins专属抵扣
    if gcoin is None:gcoin=x.get("couponPrice","")  # 存量数据用总抵扣
    for j,v in enumerate([str(i),x["name"],x["address"],str(gcoin),str(x.get("discountPrice","")),str(x.get("price",""))]):
        frun(cells[j].paragraphs[0]).text=v
final=os.path.join(OUT,f"全国GCoins_{len(done)}城_{datetime.now():%Y%m%d_%H%M%S}.docx")
doc.save(final)
print(f"\n✅ {len(done)}城 {len(all_res)}家 → {final}")
import subprocess;subprocess.Popen(["open",final])
