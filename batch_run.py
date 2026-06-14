# -*- coding: utf-8 -*-
"""批量抓取 49 城 G-Coins 酒店，排序：抵扣最高→付款最低，宋体中文表格"""
import os,sys,json,re,base64,requests,urllib3;urllib3.disable_warnings()
from datetime import datetime
sys.path.insert(0,'.')
import hopegoo_api as api

OUT="/Users/congwenjie/Desktop/未命名文件夹";os.makedirs(OUT,exist_ok=True)
DB=os.path.expanduser("~/Library/Application Support/com.reqable.macosx/box/data.mdb")

# 1. 取最新 token
raw=open(DB,"rb").read();best=None
for blk in re.findall(rb'H4sI[A-Za-z0-9+/=]{80,}',raw):
    try:d=json.loads(__import__('gzip').decompress(base64.b64decode(blk+b'=='*(-len(blk)%4))).decode("utf-8","replace"))
    except:continue
    if not isinstance(d,dict):continue
    s=d.get("session",{})
    if s.get("connection",{}).get("originHost","")!="hotel.reshg.com":continue
    p=s.get("request",{}).get("requestLine",{}).get("path","")
    if "/tapi/v2/list?" not in p:continue
    hs=s.get("request",{}).get("headers",[]);low=[str(h).lower() for h in hs]
    if any(l.startswith("sectoken:") for l in low) and any(l.startswith("dun-token:") for l in low):
        ts=s.get("timestamp",0)
        if best is None or ts>best[0]:best=(ts,p,hs)
if not best:print("无token!请先App搜一次");exit()
ts,p,hs=best;age=(datetime.now().timestamp()-ts/1e6)
print(f"Token: {datetime.fromtimestamp(ts/1e6):%H:%M:%S} ({age:.0f}秒前)")
if age>120:print("⚠️ 超2分钟，请重新App搜!");exit()

# 2. 构建headers
hd={};cookies=[]
for h in hs:
    s2=str(h);i=s2.find(":")
    if i<0:continue
    nm=s2[:i].strip();v=s2[i+1:].strip()
    if nm.lower()=="cookie":cookies.append(v)
    else:hd[nm]=v
hd["Cookie"]="; ".join(cookies)

# 3. 49城
CITIES="北京 上海 广州 深圳 成都 杭州 重庆 武汉 苏州 西安 南京 长沙 郑州 天津 合肥 青岛 东莞 宁波 佛山 厦门 大连 沈阳 昆明 南昌 哈尔滨 泉州 常州 南通 烟台 温州 长春 南宁 贵阳 石家庄 太原 珠海 嘉兴 金华 绍兴 潍坊 徐州 惠州 台州 呼和浩特 乌鲁木齐 扬州 中山 保定 兰州".split()
cmap=api.load_city_map()
s=requests.Session();s.verify=False
all=[]

for i,name in enumerate(CITIES):
    cid=cmap.get(name)
    if not cid:print(f"[{i+1}/{len(CITIES)}] ⚠️ {name}");continue
    r=s.get("https://hotel.reshg.com/tapi/v2/list",params={"city":str(cid),"filterList":"8888_4,391010_1","inDate":"2026-06-14","outDate":"2026-06-15","pageIndex":"0","pageSize":"20","adultsNumber":"1","currency":"CNY","ref":"hopegooh5","needTitleBar":"0","scriptVersion":"0.2.15"},headers=hd,timeout=15)
    d=r.json();hl=(d.get("data")or{}).get("hotelList")or[]
    for hh in hl:
        all.append({"name":(hh.get("hotelName")or"").strip(),"address":(hh.get("hotelAddress")or"").strip()or"-",
            "couponPrice":hh.get("couponPrice",""),"discountPrice":hh.get("discountPrice",""),"price":hh.get("price",""),"_city":name})
    ec=d.get("errorCode","")
    print(f"[{i+1}/{len(CITIES)}] {name}: {len(hl)}家 累计{len(all)}"+(f" 🔴token过期!" if str(ec)=="-99" else ""))
    if str(ec)=="-99":break

# 4. 排序：抵扣最高→付款最低
all.sort(key=lambda r:(-(r.get("couponPrice")or 0),r.get("discountPrice")or 1e9))

# 5. 出宋体中文表格文档
from docx import Document;from docx.shared import Pt;from docx.enum.text import WD_ALIGN_PARAGRAPH;from docx.oxml.ns import qn

def font_run(p,bold=False,size=9):
    rr=p.add_run("");rr.bold=bold;rr.font.size=Pt(size);rr.font.name='宋体'
    rr._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    return rr

doc=Document()
style=doc.styles['Normal'];style.font.name='宋体';style.font.size=Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

t=doc.add_paragraph();t.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=font_run(t,True,18);fr.text="HopeGoo G-Coins 酒店筛选报告"

info=doc.add_paragraph()
for l in [f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}","目标：全国49城","日期：2026-06-14 至 2026-06-15",
           f"排序：抵扣金额从高到低，付款金额从低到高","酒店数：{len(all)} 家"]:
    fr=font_run(info,False,10);fr.text=l+"\n"

cols=["序号","酒店名称","酒店地址","抵扣金额(CNY)","付款金额(CNY)","房价"]
tb=doc.add_table(rows=1,cols=6);tb.style="Light Grid Accent 1"
for i,c in enumerate(cols):
    fr=font_run(tb.rows[0].cells[i].paragraphs[0],True,10);fr.text=c

for i,x in enumerate(all,1):
    cells=tb.add_row().cells
    for j,v in enumerate([str(i),x["name"],x["address"],str(x.get("couponPrice","")),str(x.get("discountPrice","")),str(x.get("price",""))]):
        fr=font_run(cells[j].paragraphs[0],False,9);fr.text=v

final=os.path.join(OUT,f"全国GCoins_{datetime.now():%Y%m%d_%H%M%S}.docx")
doc.save(final)
print(f"\n✅ {len(all)}家 | 抵扣最高→付款最低 | 宋体 | 📄{final}")

import subprocess
subprocess.Popen(["open",final])
