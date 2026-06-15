# -*- coding: utf-8 -*-
"""HopeGoo G-Coins 酒店筛选（App 接口版）—— 粘贴地区文本自动识别，抵扣最高优先"""
import os,sys,csv,queue,threading,subprocess,re,socket,time
import tkinter as tk
from tkinter import ttk,filedialog,messagebox,scrolledtext
from datetime import datetime,timedelta
import hopegoo_api as api

LOG_Q=queue.Queue()
def log(m):LOG_Q.put(f"[{datetime.now():%H:%M:%S}] {m}")
def open_file(path):
    if sys.platform=="darwin":subprocess.Popen(["open",path])
    elif sys.platform.startswith("win"):os.startfile(path)
    else:subprocess.Popen(["xdg-open",path])

def t2s(text):
    """繁体中文→简体中文（优先用 opencc，不可用时用内置映射）"""
    if not text:return text
    try:
        r=subprocess.run(["opencc","-c","t2s"],input=str(text),capture_output=True,text=True,timeout=5)
        if r.returncode==0:return r.stdout.strip()
    except:pass
    try:
        from t2s_lite import t2s as _t2s;return _t2s(text)
    except:return str(text)

def write_docx(res,path,h):
    from docx import Document;from docx.shared import Pt;from docx.enum.text import WD_ALIGN_PARAGRAPH;from docx.oxml.ns import qn
    def frun(p,bold=False,size=9):
        r=p.add_run("");r.bold=bold;r.font.size=Pt(size);r.font.name='宋体';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');return r
    doc=Document()
    s=doc.styles['Normal'];s.font.name='宋体';s.font.size=Pt(10);s.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    t=doc.add_paragraph();t.alignment=WD_ALIGN_PARAGRAPH.CENTER;frun(t,True,18).text="HopeGoo G-Coins 酒店筛选报告"
    info=doc.add_paragraph()
    info_lines=[f"生成时间：{h['ts']}",f"目标地区：{h['city']}",f"入住：{h['in']}  离店：{h['out']}  人数：{h['adults']}",
                f"筛选：付款 < {h.get('max_pay','?')}元  且  G-Coins抵扣 > {h.get('min_coupon','?')}元",
                "排序：G-Coins抵扣从高到低",f"酒店数：{h['n']} 家"]
    for l in info_lines:
        frun(info,False,10).text=l+"\n"
    cols=["序号","酒店名称","酒店地址","G-Coins抵扣(CNY)","付款金额(CNY)","房价"]
    tb=doc.add_table(rows=1,cols=6);tb.style="Light Grid Accent 1"
    for i,c in enumerate(cols):frun(tb.rows[0].cells[i].paragraphs[0],True,10).text=c
    for i,x in enumerate(res,1):
        cells=tb.add_row().cells
        gcoin=x.get("gcoinAmount") or x.get("couponPrice","")
        for j,v in enumerate([str(i),t2s(x["name"]),t2s(x["address"]),str(gcoin),str(x.get("discountPrice","")),str(x.get("price",""))]):
            frun(cells[j].paragraphs[0],False,9).text=v
    if not res:frun(doc.add_paragraph(),False,10).text="（未找到 G-Coins 酒店）"
    doc.save(path)

class App:
    def __init__(self,root):
        self.root=root;root.title("HopeGoo G-Coins 酒店筛选")
        root.geometry("820x750");root.minsize(700,600)
        self.cities=api.load_city_map();self.last_path=None;self.parsed_cities=[]
        self.proxy_process=None
        self._ui();self._poll()

    def _ui(self):
        pad={"padx":10,"pady":5}
        # 地区粘贴框
        cf=ttk.LabelFrame(self.root,text="地区（粘贴文本，自动识别城市名，也支持逐行输入）")
        cf.pack(fill="x",**pad)
        self.city_text=tk.Text(cf,height=6,font=("Menlo",13))
        self.city_text.pack(fill="x",padx=6,pady=6)
        btn_row=ttk.Frame(cf);btn_row.pack(fill="x",padx=6,pady=2)
        ttk.Label(btn_row,text=f"已知 {len(self.cities)} 城，支持简体/繁体").pack(side="left")
        ttk.Button(btn_row,text="🔍 识别地区",command=self.on_parse).pack(side="right",padx=4)
        self.parse_label=tk.StringVar(value="")
        ttk.Label(btn_row,textvariable=self.parse_label,foreground="#007AFF").pack(side="right",padx=8)

        # 日期
        df=ttk.LabelFrame(self.root,text="日期")
        df.pack(fill="x",**pad)
        today=datetime.now()
        ttk.Label(df,text="入住：").grid(row=0,column=0,sticky="w",padx=8,pady=6)
        self.ind=tk.StringVar(value=today.strftime("%Y-%m-%d"))
        ttk.Entry(df,textvariable=self.ind,width=14).grid(row=0,column=1,sticky="w",padx=4)
        ttk.Label(df,text="离店：").grid(row=0,column=2,sticky="w",padx=8)
        self.outd=tk.StringVar(value=(today+timedelta(days=1)).strftime("%Y-%m-%d"))
        ttk.Entry(df,textvariable=self.outd,width=14).grid(row=0,column=3,sticky="w",padx=4)
        ttk.Label(df,text="人数：").grid(row=0,column=4,sticky="w",padx=8)
        self.adults=tk.StringVar(value="1")
        ttk.Spinbox(df,from_=1,to=10,textvariable=self.adults,width=4).grid(row=0,column=5,sticky="w")

        # 筛选+输出
        of=ttk.LabelFrame(self.root,text="筛选条件 & 输出")
        of.pack(fill="x",**pad)
        ttk.Label(of,text="付款 <").grid(row=0,column=0,sticky="w",padx=4)
        self.max_pay=tk.StringVar(value="30")
        ttk.Entry(of,textvariable=self.max_pay,width=6).grid(row=0,column=1,sticky="w")
        ttk.Label(of,text="元  且  抵扣 >").grid(row=0,column=2,sticky="w")
        self.min_coupon=tk.StringVar(value="10")
        ttk.Entry(of,textvariable=self.min_coupon,width=6).grid(row=0,column=3,sticky="w")
        ttk.Label(of,text="元").grid(row=0,column=4,sticky="w")
        ttk.Label(of,text="格式:").grid(row=0,column=5,sticky="e",padx=4)
        self.fmt=tk.StringVar(value="表格(Word)")
        ttk.Combobox(of,textvariable=self.fmt,values=["表格(Word)","TXT","CSV"],width=10,state="readonly").grid(row=0,column=6,sticky="w")
        ttk.Label(of,text="输出:").grid(row=1,column=0,sticky="w",padx=4)
        self.outdir=tk.StringVar(value=os.path.join(os.path.dirname(os.path.abspath(__file__)),"output"))
        ttk.Entry(of,textvariable=self.outdir,width=42).grid(row=1,column=1,columnspan=5,sticky="we",padx=4)
        ttk.Button(of,text="📂",width=3,command=self.on_dir).grid(row=1,column=6)
        of.columnconfigure(3,weight=1)

        # 按钮
        bar=ttk.Frame(self.root);bar.pack(fill="x",**pad)

        # 代理控制
        proxy_frm=ttk.LabelFrame(bar,text="内置代理")
        proxy_frm.pack(side="left",padx=4)
        self.proxy_btn=ttk.Button(proxy_frm,text="🔌 启动代理",command=self.on_proxy)
        self.proxy_btn.pack(side="left",padx=2)
        self.proxy_status=tk.StringVar(value="⚪ 未启动")
        ttk.Label(proxy_frm,textvariable=self.proxy_status,width=8).pack(side="left",padx=2)
        ip=self.get_lan_ip()
        ttk.Label(proxy_frm,text=f"本机IP: {ip}",foreground="#666").pack(side="left",padx=4)
        ttk.Button(proxy_frm,text="📁 加载Token文件",command=self.on_load_token,width=14).pack(side="left",padx=4)

        self.start_btn=ttk.Button(bar,text="▶ 开始筛选",command=self.on_start);self.start_btn.pack(side="left",padx=4)
        ttk.Button(bar,text="🏙 刷新城市表",command=self.on_cities).pack(side="left",padx=4)
        self.open_btn=ttk.Button(bar,text="📄 打开文件",command=self.on_open,state="disabled");self.open_btn.pack(side="left",padx=4)

        # 日志
        lf=ttk.LabelFrame(self.root,text="日志");lf.pack(fill="both",expand=True,**pad)
        self.logbox=scrolledtext.ScrolledText(lf,wrap="word",height=12,font=("Menlo",11))
        self.logbox.pack(fill="both",expand=True,padx=6,pady=6)
        self.status=tk.StringVar(value="就绪")
        ttk.Label(self.root,textvariable=self.status,anchor="w",relief="sunken").pack(fill="x")

    def on_dir(self):
        d=filedialog.askdirectory(initialdir=self.outdir.get()or".");d and self.outdir.set(d)

    def on_load_token(self):
        """手动加载 reshg_req.json token 文件（免代理模式）"""
        f=filedialog.askopenfilename(title="选择 reshg_req.json",filetypes=[("JSON文件","*.json")])
        if not f:return
        try:
            import shutil
            # EXE同目录
            target=os.path.join(os.path.dirname(sys.executable) if getattr(sys,'frozen',False) else os.path.dirname(os.path.abspath(__file__)),"reshg_req.json")
            shutil.copy(f,target)
            # 也更新 api 模块的路径
            api.REQ_TEMPLATE=target
            log(f"✅ Token文件已加载: {os.path.basename(f)}")
            messagebox.showinfo("加载成功","Token文件已就绪，可直接开始筛选")
        except Exception as e:
            messagebox.showerror("加载失败",str(e))

    def get_lan_ip(self):
        try:
            s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM)
            s.connect(('10.255.255.255',1));ip=s.getsockname()[0];s.close();return ip
        except:return "0.0.0.0"

    def on_proxy(self):
        if self.proxy_process and self.proxy_process.poll() is None:
            self.proxy_process.terminate();self.proxy_process=None
            self.proxy_btn.config(text="🔌 启动代理");self.proxy_status.set("⚪ 已停止")
            log("代理已停止")
            return
        try:
            addon=os.path.join(os.path.dirname(os.path.abspath(__file__)),"capture_addon.py")
            # 兼容打包后的路径
            if getattr(sys,'frozen',False):
                addon=os.path.join(sys._MEIPASS,"capture_addon.py")

            # 尝试多种方式启动 mitmproxy
            started=False
            mitm_bin=None
            # 1) EXE 内打包的 mitmdump
            if getattr(sys,'frozen',False):
                bundled=os.path.join(sys._MEIPASS,"mitmdump.exe")
                if os.path.exists(bundled): mitm_bin=bundled
            # 2) 系统 PATH 里的 mitmdump
            if not mitm_bin:
                for p in ["mitmdump","mitmdump.exe"]:
                    if subprocess.run(["where",p],capture_output=True,shell=True).returncode==0:
                        mitm_bin=p;break
            # 3) Python 模块方式
            cmds=[]
            if mitm_bin: cmds.append([mitm_bin,"-s",addon,"--listen-port","8080","--set","block_global=false"])
            cmds.append([sys.executable,"-m","mitmproxy.tools._main","-s",addon,"--listen-port","8080","--set","block_global=false"])
            for cmd in cmds:
                try:
                    self.proxy_process=subprocess.Popen(cmd,stdout=subprocess.DEVNULL,stderr=subprocess.DEVNULL)
                    started=True;log(f"mitmproxy 启动成功")
                    break
                except: continue
            if not started: raise RuntimeError("无法启动 mitmproxy，请运行 pip install mitmproxy")

            ip=self.get_lan_ip()
            self.proxy_btn.config(text="⏹ 停止代理");self.proxy_status.set("🟢 运行中")
            log(f"代理已启动！手机WiFi代理设为 {ip}:8080")
            messagebox.showinfo("代理已启动",
                f"1. 手机WiFi代理设为 {ip}:8080\n"
                f"2. 手机浏览器打开 mitm.it 安装CA证书\n"
                f"3. App搜任意城市 → Token自动捕获\n"
                f"4. 点「开始筛选」即可")
        except Exception as e:
            err=str(e)
            msg=f"代理启动失败：{err}\n\n"
            if sys.platform=='win32':
                msg+="可能原因：缺少 Visual C++ 运行库\n\n"
                msg+="一键修复：下载安装 VC++ Redistributable\n"
                msg+="https://aka.ms/vs/17/release/vc_redist.x64.exe\n\n"
                msg+="安装后重启程序即可。"
            else:
                msg+="请运行: pip install mitmproxy"
            messagebox.showerror("启动失败",msg)
    def on_open(self):
        try:open_file(self.last_path)
        except Exception as e:messagebox.showerror("打开失败",str(e))
    def on_cities(self):
        try:self.cities=api.refresh_city_map_from_reqable(log=log);messagebox.showinfo("OK",f"已更新 {len(self.cities)} 城")
        except Exception as e:messagebox.showerror("失败",str(e))

    def on_parse(self):
        text=self.city_text.get("1.0","end-1c")
        found=[];seen=set()
        for cn in sorted(self.cities,key=lambda x:-len(x)):
            if cn in text and cn not in seen:found.append(cn);seen.add(cn)
        for m in re.finditer(r'([一-鿿]{2,4})市',text):
            n=m.group(1);c=self.cities.get(n) or self.cities.get(n.replace("市",""))
            if c and n not in seen:found.append(n);seen.add(n)
        if found:
            self.parsed_cities=found
            self.parse_label.set(f"识别到 {len(found)} 个：{'、'.join(found[:10])}{'…' if len(found)>10 else ''}")
        else:
            self.parsed_cities=[]
            self.parse_label.set("未识别到地区")

    def _resolve(self,token):
        token=token.strip()
        if not token:return None
        if token.isdigit():return int(token)
        for k in [token,token.replace("市",""),token+"市"]:
            if self.cities.get(k):return self.cities[k]
        return None

    def on_start(self):
        if not hasattr(self,'parsed_cities') or not self.parsed_cities:self.on_parse()
        items=self.parsed_cities
        if not items:
            text=self.city_text.get("1.0","end-1c")
            items=[l.strip() for l in text.replace("，",",").replace("、","\n").replace(" ","\n").split("\n") if l.strip()]
        if not items:return messagebox.showerror("参数错误","请粘贴含地区名的文本")
        try:datetime.strptime(self.ind.get(),"%Y-%m-%d");datetime.strptime(self.outd.get(),"%Y-%m-%d");adults=int(self.adults.get())
        except ValueError:return messagebox.showerror("参数错误","日期/人数格式不对")
        resolved=[]
        for it in items:
            cid=self._resolve(it)
            if cid is None:
                if not messagebox.askyesno("未知地区",f"找不到「{it}」的城市ID。跳过？"):return
            else:resolved.append((it,cid))
        if not resolved:return
        if not os.path.exists(api.REQ_TEMPLATE):
            if messagebox.askyesno("缺少Token","未检测到Token文件。\n\n点「是」= 加载已有的 reshg_req.json\n点「否」= 取消"):
                self.on_load_token()
            if not os.path.exists(api.REQ_TEMPLATE):return
        self.start_btn.config(state="disabled");self.status.set("抓取中…")
        threading.Thread(target=self._worker,args=(resolved,adults),daemon=True).start()

    def _worker(self,resolved,adults):
        try:
            all_res=[];i=0
            while i<len(resolved):
                name,cid=resolved[i]
                log(f"=== {name} ===")
                res=api.scrape_city(cid,self.ind.get(),self.outd.get(),adults=adults,log=log)
                # 筛选：付款 < max_pay, 抵扣 > min_coupon
                max_p=float(self.max_pay.get() or "30")
                min_c=float(self.min_coupon.get() or "10")
                raw_count=len(res)
                res=[r for r in res if (r.get("discountPrice")or 1e9)<max_p and (r.get("gcoinAmount")or r.get("couponPrice")or 0)>min_c]
                filtered=raw_count-len(res)
                if filtered>0:log(f"  筛选: {raw_count}家→{len(res)}家 (付款<{max_p}元, 抵扣>{min_c}元)")
                if not res and raw_count==0:
                    log("  token过期，等待App搜索…")
                    self.root.after(0,lambda:self.status.set("Token过期！"))
                    choice=tk.StringVar(value="")
                    dlg=tk.Toplevel(self.root);dlg.title("Token 过期")
                    dlg.geometry("300x150")
                    ttk.Label(dlg,text="Token已过期，请选择：",padding=15).pack()
                    btn_frm=ttk.Frame(dlg);btn_frm.pack(pady=10)
                    def do_wait():choice.set("wait");dlg.destroy()
                    def do_end():choice.set("end");dlg.destroy()
                    ttk.Button(btn_frm,text="继续等待（App搜后自动续）",command=do_wait).pack(side="left",padx=5)
                    ttk.Button(btn_frm,text="结束并生成文档",command=do_end).pack(side="left",padx=5)
                    self.root.wait_window(dlg)
                    if choice.get()=="wait":
                        log("  等待新token…请在App搜索/切换账号")
                        for _ in range(60):
                            time.sleep(2)
                            if os.path.exists(api.REQ_TEMPLATE):
                                try:
                                    tpl=api.load_template()
                                    # 测试token是否有效
                                    break
                                except:pass
                        log("  继续尝试…")
                        continue  # 重试当前城市
                    else:
                        log("  用户选择结束，生成当前文档")
                        break  # 跳出循环，生成文档
                for r in res:r["_city_label"]=name
                all_res+=res
                log(f"  → {len(res)} 家 (累计{len(all_res)})")
                i+=1
            # 排序：抵扣最高→付款最低
            all_res.sort(key=lambda r:(-(r.get("gcoinAmount")or r.get("couponPrice")or 0),(r.get("discountPrice")or 1e9)))
            done=sum(1 for name,_ in resolved if any(r.get("_city_label")==name for r in all_res))
            h={"ts":datetime.now().strftime("%Y-%m-%d %H:%M:%S"),"city":f"{done}个地区","in":self.ind.get(),"out":self.outd.get(),"adults":adults,"n":len(all_res),
               "max_pay":self.max_pay.get(),"min_coupon":self.min_coupon.get()}
            os.makedirs(self.outdir.get(),exist_ok=True)
            safe="".join(c for c in f"GCoins酒店{done}城" if c not in '\\/:*?"<>| ')[:30] or "酒店"
            path=os.path.join(self.outdir.get(),f"{safe}_{datetime.now():%Y-%m-%d_%H%M%S}.docx")
            write_docx(all_res,path,h)
            log(f"文档已生成：{path}")
            self.last_path=path
            self.root.after(0,lambda:self.open_btn.config(state="normal"))
            try:open_file(path)
            except:pass
            self.root.after(0,lambda:messagebox.showinfo("完成",f"共 {len(all_res)} 家。\n{path}"))
            self.root.after(0,lambda:self.status.set(f"完成，共 {len(all_res)} 家"))
        except Exception as e:
            log(f"出错：{e}")
            self.root.after(0,lambda:messagebox.showerror("错误",str(e)))
            self.root.after(0,lambda:self.status.set("失败"))
        finally:self.root.after(0,lambda:self.start_btn.config(state="normal"))

    def _poll(self):
        while True:
            try:m=LOG_Q.get_nowait()
            except queue.Empty:break
            self.logbox.insert("end",m+"\n");self.logbox.see("end")
        self.root.after(300,self._poll)

def main():
    root=tk.Tk()
    try:
        st=ttk.Style(root)
        if sys.platform=="darwin":st.theme_use("aqua")
    except:pass
    App(root);root.mainloop()

if __name__=="__main__":main()
