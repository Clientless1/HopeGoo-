# -*- coding: utf-8 -*-
"""全自动批量抓取：内置代理→监听token→自动续跑→简体中文文档"""
import os,sys,json,re,time,subprocess,socket
from datetime import datetime
sys.path.insert(0,'.')
import hopegoo_api as api
import requests,urllib3;urllib3.disable_warnings()

OUT="/Users/congwenjie/Desktop/未命名文件夹";os.makedirs(OUT,exist_ok=True)
PROGRESS=os.path.join(OUT,"progress_full.json")
TOKEN_FILE=os.path.join(os.path.dirname(os.path.abspath(__file__)),"reshg_req.json")
STATUS_FILE=os.path.join(os.path.dirname(os.path.abspath(__file__)),"token_updated.txt")
IP=socket.gethostbyname(socket.gethostname())
PROXY_PORT=8080

def log(msg):
    t=datetime.now().strftime("%H:%M:%S")
    print(f"[{t}] {msg}")

def t2s(text):
    """繁体→简体"""
    if not text:return text
    try:
        r=subprocess.run(["opencc","-c","t2s"],input=text,capture_output=True,text=True,timeout=5)
        return r.stdout.strip() if r.returncode==0 else text
    except:return text

def start_proxy():
    addon=os.path.join(os.path.dirname(os.path.abspath(__file__)),"capture_addon.py")
    return subprocess.Popen(["mitmdump","-s",addon,"--listen-port",str(PROXY_PORT),"--set","block_global=false"],
        stdout=subprocess.DEVNULL,stderr=subprocess.DEVNULL)

def wait_for_token(prev_ts=0,timeout=300):
    """等待新token出现，返回新时间戳"""
    log(f"等待App请求… 手机代理设为 {IP}:{PROXY_PORT}")
    start=time.time()
    while time.time()-start<timeout:
        if os.path.exists(STATUS_FILE):
            try:ts=float(open(STATUS_FILE).read().strip())
            except:ts=0
            if ts>prev_ts:
                age=time.time()-ts
                if age<60:
                    log(f"✅ 检测到新Token! ({age:.0f}秒前)")
                    return ts
        time.sleep(2)
        elapsed=time.time()-start
        if elapsed%30<2:print(f"  等待中… {elapsed:.0f}s",end="\r")
    return 0

def get_hd_from_token():
    tpl=json.load(open(TOKEN_FILE))
    import urllib.parse as _up
    parsed=_up.urlparse("https://hotel.reshg.com"+tpl["path"])
    base_params=dict(_up.parse_qsl(parsed.query))
    hd=dict(tpl["headers"]);hd["Cookie"]="; ".join(tpl["cookies"])
    return hd,base_params

# 城市列表
CITIES="""北京 天津 石家庄 唐山 秦皇岛 邯郸 邢台 保定 张家口 承德 沧州 廊坊 衡水
太原 大同 阳泉 长治 晋城 朔州 忻州 吕梁 晋中 临汾 运城
呼和浩特 包头 乌海 赤峰 通辽 鄂尔多斯 呼伦贝尔 巴彦淖尔 乌兰察布
沈阳 大连 鞍山 抚顺 本溪 丹东 锦州 营口 阜新 辽阳 盘锦 铁岭 朝阳 葫芦岛
长春 吉林 四平 辽源 通化 白山 松原 白城
哈尔滨 齐齐哈尔 鸡西 鹤岗 双鸭山 大庆 伊春 佳木斯 七台河 牡丹江 黑河 绥化
上海 南京 无锡 徐州 常州 苏州 南通 连云港 淮安 盐城 扬州 镇江 泰州 宿迁
杭州 宁波 温州 嘉兴 湖州 绍兴 金华 衢州 舟山 台州 丽水
合肥 芜湖 蚌埠 淮南 马鞍山 淮北 铜陵 安庆 黄山 滁州 阜阳 宿州 六安 亳州 池州 宣城
福州 厦门 莆田 三明 泉州 漳州 南平 龙岩 宁德
南昌 景德镇 萍乡 九江 新余 鹰潭 赣州 吉安 宜春 抚州 上饶
济南 青岛 淄博 枣庄 东营 烟台 潍坊 济宁 泰安 威海 日照 临沂 德州 聊城 滨州 菏泽
郑州 开封 洛阳 平顶山 安阳 鹤壁 新乡 焦作 濮阳 许昌 漯河 三门峡 南阳 商丘 信阳 周口 驻马店
武汉 黄石 十堰 宜昌 襄阳 鄂州 荆门 孝感 荆州 黄冈 咸宁 随州 仙桃 潜江 天门
长沙 株洲 湘潭 衡阳 邵阳 岳阳 常德 张家界 益阳 郴州 永州 怀化 娄底
广州 深圳 珠海 汕头 佛山 韶关 湛江 茂名 肇庆 惠州 梅州 汕尾 河源 阳江 清远 东莞 中山 潮州 揭阳 云浮
南宁 柳州 桂林 梧州 北海 防城港 钦州 贵港 玉林 百色 贺州 河池 来宾 崇左
海口 三亚 三沙 儋州
重庆 成都 自贡 攀枝花 泸州 德阳 绵阳 广元 遂宁 内江 乐山 南充 眉山 宜宾 广安 达州 雅安 巴中 资阳
贵阳 六盘水 遵义 安顺 毕节 铜仁
昆明 曲靖 玉溪 保山 昭通 丽江 普洱 临沧
拉萨 日喀则 昌都 林芝 那曲
西安 铜川 宝鸡 咸阳 渭南 延安 汉中 榆林 安康 商洛
兰州 嘉峪关 金昌 白银 天水 武威 张掖 平凉 酒泉 庆阳 定西 陇南
西宁 海东 银川 石嘴山 吴忠 固原 中卫
乌鲁木齐 克拉玛依 吐鲁番 哈密 石河子 阿拉尔 五家渠""".split()

CITY_SET=[];seen=set()
for t in CITIES:
    t=t.strip()
    if t and t not in seen:CITY_SET.append(t);seen.add(t)
cmap=api.load_city_map()
CITY_LIST=[(c,cmap.get(c) or cmap.get(c.replace("市",""))) for c in CITY_SET if cmap.get(c) or cmap.get(c.replace("市",""))]

# 恢复进度
if os.path.exists(PROGRESS):
    p=json.load(open(PROGRESS));all_res=p["results"];done=set(p["done"]);start_idx=p["idx"]
    log(f"恢复: {len(all_res)}家 {len(done)}城 从第{start_idx+1}城续")
else:all_res=[];done=set();start_idx=0

def save():json.dump({"results":all_res,"done":list(done),"idx":i},open(PROGRESS,"w"),ensure_ascii=False)

# 启动代理
proxy=start_proxy()
log(f"代理已启动: {IP}:{PROXY_PORT}")
log(f"手机WiFi代理设为 {IP}:{PROXY_PORT}，打开App搜任意城市")

# 等待首次token
last_token_ts=float(open(STATUS_FILE).read().strip()) if os.path.exists(STATUS_FILE) else 0
if not os.path.exists(TOKEN_FILE) or time.time()-last_token_ts>60:
    last_token_ts=wait_for_token(last_token_ts)
    if not last_token_ts:
        log("超时！请确认手机已设代理并搜索");proxy.terminate();sys.exit(1)

hd,base_params=get_hd_from_token()
s=requests.Session();s.verify=False
i=start_idx-1 if start_idx>0 else -1

while True:
    i+=1
    if i>=len(CITY_LIST):
        log("🎉 全部完成!");save();break
    name,cid=CITY_LIST[i]
    if name in done:continue

    try:
        base_params["city"]=str(cid);base_params["pageIndex"]="0";base_params["pageSize"]="20"
        r=s.get("https://hotel.reshg.com/tapi/v2/list?"+__import__('urllib').parse.urlencode(base_params),
            headers=hd,timeout=15)
        d=r.json();ec=str(d.get("errorCode",""))
        if ec=="-99":
            log(f"🔴 Token过期于{name}，等待新token…")
            log(f"   请在App里切换账号或重新搜索")
            save()
            # 等新token
            prev_ts=last_token_ts
            last_token_ts=wait_for_token(prev_ts)
            if not last_token_ts:
                log("超时，停止");break
            hd,base_params=get_hd_from_token()
            i-=1;continue  # 重试当前城市

        hl=(d.get("data")or{}).get("hotelList")or[]
        for hh in hl:
            gcoin=next((lb.get("amount") for lb in (hh.get("productLabelList")or[]) if lb.get("productLabelId")==345),None)
            all_res.append({
                "name":t2s((hh.get("hotelName")or"").strip()),
                "address":t2s((hh.get("hotelAddress")or"").strip()or"-"),
                "couponPrice":hh.get("couponPrice",""),
                "discountPrice":hh.get("discountPrice",""),
                "price":hh.get("price",""),
                "gcoinAmount":gcoin,"_city":t2s(name)})
        done.add(name)
        log(f"[{i+1}/{len(CITY_LIST)}] {name}: {len(hl)}家 (累计{len(all_res)}/{len(done)}城)")
    except Exception as e:
        log(f"[{i+1}] {name}: {e}")
        time.sleep(2)

    if (i+1)%5==0:save()

# 排序导出
all_res.sort(key=lambda r:(-(r.get("gcoinAmount")or 0),(r.get("discountPrice")or 1e9)))

from docx import Document;from docx.shared import Pt;from docx.enum.text import WD_ALIGN_PARAGRAPH;from docx.oxml.ns import qn
def frun(p,b=False,s=9):
    r=p.add_run("");r.bold=b;r.font.size=Pt(s);r.font.name='宋体';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');return r

doc=Document()
st=doc.styles['Normal'];st.font.name='宋体';st.font.size=Pt(10);st.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
t=doc.add_paragraph();t.alignment=WD_ALIGN_PARAGRAPH.CENTER;frun(t,True,18).text="HopeGoo G-Coins 全国酒店筛选报告"
info=doc.add_paragraph()
for l in [f"生成时间：{datetime.now():%Y-%m-%d %H:%M}",f"覆盖城市：{len(done)}个",f"酒店总数：{len(all_res)}家","排序：G-Coins抵扣从高到低，付款金额从低到高"]:
    frun(info).text=l+"\n"

tb=doc.add_table(rows=1,cols=6);tb.style="Light Grid Accent 1"
for i,c in enumerate(["序号","酒店名称","酒店地址","G-Coins抵扣(CNY)","付款金额(CNY)","房价"]):
    frun(tb.rows[0].cells[i].paragraphs[0],True,10).text=c
for i,x in enumerate(all_res,1):
    cells=tb.add_row().cells
    gcoin=x.get("gcoinAmount");gcoin=str(gcoin) if gcoin is not None else "-"
    for j,v in enumerate([str(i),x["name"],x["address"],gcoin,str(x.get("discountPrice","")),str(x.get("price",""))]):
        frun(cells[j].paragraphs[0]).text=v

path=os.path.join(OUT,f"全国GCoins_{len(done)}城_{datetime.now():%Y%m%d_%H%M%S}.docx")
doc.save(path)
log(f"✅ {path}")
subprocess.Popen(["open",path])
proxy.terminate()
