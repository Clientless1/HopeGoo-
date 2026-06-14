# -*- coding: utf-8 -*-
"""
HopeGoo 酒店筛选工具
----------------------------------
功能：
  1. GUI 输入：账号 / 密码 / 目标地区 / 入住日期 / 离店日期 / 入住人数 / 优惠上限
  2. 通过 Playwright 自动打开 hopegoo，登录并把货币切换为 CNY
  3. 进入酒店列表页，输入地区并按“价格”排序（低价优先）
  4. 滚动加载，逐家读取房价
  5. 计算【付款金额 = 房价 − 优惠券(默认20)】，保留“付款金额 < 上限(默认25)”的酒店
  6. 输出为正式文档（DOCX / TXT / CSV）

判定规则（与用户确认）：
  - 抵扣金额 = 优惠券面值（默认 CNY20）
  - 付款金额 = 房价 − 优惠券 = 最终要付的钱
  - 只要【付款金额 < 上限(默认25)】即算符合（价格达标就算，无需逐家进详情页）

依赖：
  pip install playwright python-docx
  playwright install chromium      # 若自带 chromium 损坏，本工具会自动改用系统 Chrome
运行：
  python hopegoo_hotel_filter.py
"""

import os
import re
import sys
import csv
import time
import queue
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from datetime import datetime, timedelta

# ============================================================
# 繁简转换（确保输出文档为中文简体）
# ============================================================
def to_simplified(text: str) -> str:
    """将繁体中文转换为简体中文。"""
    if not text:
        return text
    try:
        from opencc import OpenCC
        cc = OpenCC('t2s')
        return cc.convert(text)
    except ImportError:
        # 回退到内置映射表（常用字）
        _SIMPLIFIED_MAP = {
            '為': '为', '為': '为', '為': '为', '為': '为', '為': '为',
            '為': '为', '為': '为', '為': '为', '為': '为', '為': '为',
            '機': '机', '電': '电', '話': '话', '號': '号', '碼': '码',
            '價': '价', '錢': '钱', '國': '国', '際': '际', '區': '区',
            '區': '区', '發': '发', '達': '达', '開': '开', '關': '关',
            '關': '关', '係': '系', '經': '经', '營': '营', '業': '业',
            '業': '业', '員': '员', '會': '会', '計': '计', '畫': '画',
            '畫': '画', '圖': '图', '書': '书', '報': '报', '紙': '纸',
            '館': '馆', '飯': '饭', '店': '店', '旅': '旅', '館': '馆',
            '車': '车', '站': '站', '路': '路', '街': '街', '巷': '巷',
            '門': '门', '窗': '窗', '牆': '墙', '樓': '楼', '房': '房',
            '間': '间', '衛': '卫', '生': '生', '廁': '厕', '所': '所',
            '廚': '厨', '房': '房', '廳': '厅', '客': '客', '廳': '厅',
            '室': '室', '臥': '卧', '室': '室', '浴': '浴', '室': '室',
            '陽': '阳', '臺': '台', '花': '花', '園': '园', '草': '草',
            '地': '地', '坪': '坪', '樹': '树', '木': '木', '林': '林',
            '山': '山', '水': '水', '河': '河', '海': '海', '江': '江',
            '湖': '湖', '島': '岛', '岸': '岸', '港': '港', '灣': '湾',
            '道': '道', '路': '路', '街': '街', '巷': '巷', '弄': '弄',
            '里': '里', '村': '村', '鎮': '镇', '市': '市', '縣': '县',
            '省': '省', '區': '区', '州': '州', '府': '府', '郡': '郡',
            '國': '国', '界': '界', '外': '外', '內': '内', '中': '中',
            '東': '东', '西': '西', '南': '南', '北': '北', '前': '前',
            '後': '后', '左': '左', '右': '右', '上': '上', '下': '下',
            '大': '大', '小': '小', '多': '多', '少': '少', '長': '长',
            '長': '长', '短': '短', '高': '高', '低': '低', '寬': '宽',
            '寬': '宽', '窄': '窄', '厚': '厚', '薄': '薄', '深': '深',
            '淺': '浅', '重': '重', '輕': '轻', '輕': '轻', '快': '快',
            '慢': '慢', '遠': '远', '近': '近', '舊': '旧', '新': '新',
            '好': '好', '壞': '坏', '壞': '坏', '美': '美', '醜': '丑',
            '真': '真', '假': '假', '善': '善', '惡': '恶', '惡': '恶',
            '是': '是', '非': '非', '有': '有', '無': '无', '無': '无',
            '不': '不', '能': '能', '可': '可', '會': '会', '要': '要',
            '想': '想', '做': '做', '說': '说', '說': '说', '看': '看',
            '聽': '听', '聽': '听', '吃': '吃', '喝': '喝', '睡': '睡',
            '走': '走', '跑': '跑', '跳': '跳', '飛': '飞', '飛': '飞',
            '開': '开', '關': '关', '關': '关', '拿': '拿', '放': '放',
            '給': '给', '給': '给', '取': '取', '送': '送', '買': '买',
            '賣': '卖', '賣': '卖', '借': '借', '還': '还', '還': '还',
            '用': '用', '愛': '爱', '恨': '恨', '喜': '喜', '怒': '怒',
            '哀': '哀', '樂': '乐', '樂': '乐', '哭': '哭', '笑': '笑',
            '生': '生', '死': '死', '病': '病', '老': '老', '年': '年',
            '日': '日', '月': '月', '天': '天', '時': '时', '時': '时',
            '分': '分', '秒': '秒', '鐘': '钟', '鐘': '钟', '表': '表',
            '錢': '钱', '幣': '币', '元': '元', '角': '角', '分': '分',
            '數': '数', '字': '字', '計': '计', '算': '算', '加': '加',
            '減': '减', '減': '减', '乘': '乘', '除': '除', '等': '等',
            '於': '于', '在': '在', '和': '和', '與': '与', '及': '及',
            '或': '或', '但': '但', '而': '而', '所以': '所以', '因為': '因为',
            '如果': '如果', '雖然': '虽然', '但是': '但是', '而且': '而且',
            '只有': '只有', '只要': '只要', '必須': '必须', '必須': '必须',
            '可以': '可以', '不能': '不能', '不要': '不要', '应该': '应该',
            '能够': '能够', '可能': '可能', '需要': '需要', '应该': '应该',
            '希望': '希望', '想': '想', '要': '要', '会': '会', '不会': '不会',
            '能': '能', '不能': '不能', '可以': '可以', '不可以': '不可以',
            '好': '好', '不好': '不好', '快': '快', '慢': '慢', '多': '多',
            '少': '少', '大': '大', '小': '小', '高': '高', '低': '低',
            '長': '长', '長': '长', '短': '短', '宽': '宽', '窄': '窄',
            '深': '深', '浅': '浅', '厚': '厚', '薄': '薄', '重': '重',
            '輕': '轻', '輕': '轻', '远': '远', '近': '近', '旧': '旧',
            '新': '新', '美': '美', '丑': '丑', '真': '真', '假': '假',
            '善': '善', '惡': '恶', '惡': '恶', '对': '对', '错': '错',
            '是': '是', '非': '非', '有': '有', '無': '无', '無': '无',
            '不': '不', '很': '很', '也': '也', '都': '都', '就': '就',
            '还': '还', '再': '再', '又': '又', '更': '更', '最': '最',
            '太': '太', '非常': '非常', '特别': '特别', '十分': '十分',
            '非常': '非常', '很': '很', '挺': '挺', '比较': '比较', '稍微': '稍微',
            '一点': '一点', '一些': '一些', '许多': '许多', '很多': '很多',
            '不少': '不少', '少量': '少量', '全部': '全部', '所有': '所有',
            '任何': '任何', '每': '每', '各': '各', '其他': '其他', '另外': '另外',
            '自己': '自己', '别人': '别人', '大家': '大家', '人们': '人们',
            '他们': '他们', '她们': '她们', '它们': '它们', '我们': '我们',
            '你们': '你们', '谁': '谁', '什么': '什么', '哪里': '哪里',
            '怎么': '怎么', '为什么': '为什么', '何时': '何时', '如何': '如何',
            '多少': '多少', '几': '几', '哪': '哪', '谁': '谁', '什么': '什么',
            '这': '这', '那': '那', '这些': '这些', '那些': '那些', '这个': '这个',
            '那个': '那个', '这样': '这样', '那样': '那样', '这里': '这里',
            '那里': '那里', '这么': '这么', '那么': '那么', '如此': '如此',
            '一样': '一样', '不同': '不同', '相同': '相同', '类似': '类似',
            '好像': '好像', '似乎': '似乎', '仿佛': '仿佛', '如同': '如同',
            '比如': '比如', '例如': '例如', '等等': '等等', '以及': '以及',
            '包括': '包括', '还有': '还有', '另外': '另外', '除了': '除了',
            '但是': '但是', '可是': '可是', '然而': '然而', '不过': '不过',
            '却': '却', '反而': '反而', '虽然': '虽然', '尽管': '尽管',
            '即使': '即使', '如果': '如果', '假如': '假如', '要是': '要是',
            '万一': '万一', '只要': '只要', '只有': '只有', '除非': '除非',
            '否则': '否则', '不管': '不管', '无论': '无论', '不论': '不论',
            '因为': '因为', '由于': '由于', '所以': '所以', '因此': '因此',
            '于是': '于是', '从而': '从而', '以致': '以致', '既然': '既然',
            '就': '就', '便': '便', '才': '才', '又': '又', '再': '再',
            '也': '也', '还': '还', '更': '更', '甚至': '甚至', '连': '连',
            '都': '都', '也': '也', '又': '又', '再': '再', '仍然': '仍然',
            '依然': '依然', '还是': '还是', '始终': '始终', '一直': '一直',
            '已经': '已经', '曾经': '曾经', '刚刚': '刚刚', '正在': '正在',
            '将要': '将要', '就要': '就要', '快要': '快要', '马上': '马上',
            '立刻': '立刻', '立即': '立即', '忽然': '忽然', '突然': '突然',
            '渐渐': '渐渐', '慢慢': '慢慢', '逐渐': '逐渐', '终于': '终于',
            '最后': '最后', '首先': '首先', '其次': '其次', '然后': '然后',
            '接着': '接着', '随后': '随后', '之后': '之后', '之前': '之前',
            '以来': '以来', '至今': '至今', '暂时': '暂时', '临时': '临时',
            '永远': '永远', '永久': '永久', '偶尔': '偶尔', '经常': '经常',
            '常常': '常常', '总是': '总是', '始终': '始终', '一直': '一直',
            '从来': '从来', '向来': '向来', '一贯': '一贯', '从来': '从来',
            '从不': '从不', '几乎': '几乎', '差不多': '差不多', '大概': '大概',
            '大约': '大约', '左右': '左右', '上下': '上下', '前后': '前后',
            '之间': '之间', '以上': '以上', '以下': '以下', '以内': '以内',
            '以外': '以外', '以前': '以前', '以后': '以后', '之前': '之前',
            '之后': '之后', '当中': '当中', '中间': '中间', '旁边': '旁边',
            '附近': '附近', '周围': '周围', '到处': '到处', '各处': '各处',
            '任何地方': '任何地方', '某个地方': '某个地方', '某些地方': '某些地方',
            '每处': '每处', '各处': '各处', '所有地方': '所有地方', '到处': '到处',
        }
        return ''.join(_SIMPLIFIED_MAP.get(c, c) for c in text)

# ============================================================
# 全局配置
# ============================================================
BASE_URL = "https://www.hopegoo.com"
HOME_URL = "https://www.hopegoo.com/zh-HK"
SEARCH_URL_TEMPLATE = (
    "https://www.hopegoo.com/zh-HK/hotel/hotellist"
    "?inDate={in_date}&outDate={out_date}&adultsNumber={adults}&currency=CNY"
)

PAY_THRESHOLD_DEFAULT = 25        # 最终付款上限（CNY），付款金额 < 该值算符合
COUPON_DEFAULT = 20               # 优惠券面值（CNY），用于计算付款金额 = 房价 − 优惠券
HEADLESS_DEFAULT = False          # 默认显示浏览器窗口（便于观察登录/验证码）
MAX_SCROLL_DEFAULT = 30          # 默认最多加载批次（每批约 20 家；按价格升序，超过上限会自动早停）

# 已保存账号（明文存于脚本同目录 accounts.json，便于下拉切换）
ACCOUNTS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "accounts.json")

# 启动浏览器优先用系统真实浏览器（自带 chromium 的 headless_shell 在部分机器损坏）
BROWSER_CHANNELS = ("chrome", "msedge")

LOG_QUEUE = queue.Queue()


# ============================================================
# 日志工具
# ============================================================
def log(msg: str) -> None:
    ts = datetime.now().strftime("%H:%M:%S")
    LOG_QUEUE.put(f"[{ts}] {msg}")


def open_file(path: str) -> None:
    """用系统默认程序打开文件。"""
    import subprocess
    if not path or not os.path.exists(path):
        raise FileNotFoundError(path)
    if sys.platform == "darwin":
        subprocess.Popen(["open", path])
    elif sys.platform.startswith("win"):
        os.startfile(path)  # type: ignore[attr-defined]
    else:
        subprocess.Popen(["xdg-open", path])


def load_accounts() -> list:
    """读取已保存账号列表 [{'account','password'}]。"""
    import json
    try:
        with open(ACCOUNTS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return [a for a in data if a.get("account")]
    except (FileNotFoundError, ValueError, OSError):
        pass
    return []


def save_accounts(accounts: list) -> None:
    import json
    with open(ACCOUNTS_FILE, "w", encoding="utf-8") as f:
        json.dump(accounts, f, ensure_ascii=False, indent=2)


# ============================================================
# 折扣金额解析
# ============================================================
DISCOUNT_PATTERNS = [
    re.compile(r"CNY\s*([\d,.]+)\s*折扣"),
    re.compile(r"([\d,.]+)\s*折扣"),
    re.compile(r"立減\s*CNY\s*([\d,.]+)"),
    re.compile(r"已抵扣\s*CNY\s*([\d,.]+)"),
]


def parse_discount(tags_text: str):
    """
    从卡片标签文本里解析“优惠/折扣金额”（CNY）。
    例：'訂房享優惠價 CNY 24 折扣' -> 24.0 ；'平台優惠 CNY 4 折扣' -> 4.0
    解析不到返回 None。
    """
    if not tags_text:
        return None
    t = tags_text.replace("\n", " ")
    for pat in DISCOUNT_PATTERNS:
        m = pat.search(t)
        if m:
            try:
                return float(m.group(1).replace(",", ""))
            except ValueError:
                continue
    return None


def parse_price(price_text: str):
    """从 '連稅 CNY307' 解析出 307。解析不到返回 None。"""
    if not price_text:
        return None
    m = re.search(r"CNY\s*([\d,.]+)", price_text.replace(" ", ""))
    if m:
        try:
            return float(m.group(1).replace(",", ""))
        except ValueError:
            return None
    return None


# ============================================================
# 浏览器/登录/货币 辅助
# ============================================================
def _goto_resilient(page, url, log_fn, attempts: int = 3):
    """首页/列表页较重，domcontentloaded 偶发超时；重试且超时也不致命。"""
    last = None
    for i in range(attempts):
        try:
            page.goto(url, wait_until="commit", timeout=90000)
            return True
        except Exception as e:
            last = e
            log_fn(f"  打开页面第 {i+1} 次超时，重试…")
            time.sleep(2)
    log_fn(f"  页面导航多次超时（继续尝试操作）：{last}")
    return False


def launch_browser(p, headless: bool, log_fn):
    """优先系统 Chrome/Edge，失败再回退自带 chromium。"""
    launch_kwargs = {"headless": headless, "args": ["--start-maximized"]}
    for channel in BROWSER_CHANNELS:
        try:
            log_fn(f"启动浏览器: channel={channel} headless={headless}")
            return p.chromium.launch(channel=channel, **launch_kwargs)
        except Exception as exc:
            log_fn(f"  channel={channel} 不可用: {exc}")
    log_fn("回退使用 Playwright 自带 Chromium")
    return p.chromium.launch(**launch_kwargs)


def do_login(page, account: str, password: str, log_fn) -> bool:
    """
    email-first 弹窗登录：
      点“登入” -> 填邮箱 -> 繼續 -> 用密碼登入 -> 填密码 -> 登入
    成功返回 True。
    """
    try:
        _goto_resilient(page, HOME_URL, log_fn)
        # 等首页“登入”入口出现（比 domcontentloaded 更可靠）
        page.wait_for_selector("text=登入", timeout=30000)
        time.sleep(1.5)

        # 已登录则直接返回
        if "登入 / 註冊" not in page.evaluate("() => document.body.innerText"):
            log_fn("已是登录状态")
            return True

        # 打开登录弹窗（弹窗有动画，重试几次直到邮箱框出现）
        email = None
        for attempt in range(3):
            try:
                page.locator("text=登入").first.click(timeout=8000)
            except Exception:
                pass
            try:
                email = page.locator("input[name=email]")
                email.first.wait_for(state="visible", timeout=8000)
                break
            except Exception:
                log_fn(f"  登录弹窗未就绪，重试 {attempt+1}")
                time.sleep(1.5)
        if email is None:
            log_fn("  无法打开登录弹窗")
            return False

        # 邮箱
        email.first.click()
        email.first.fill(account, timeout=10000)
        time.sleep(0.6)
        page.locator(":text('繼續')").first.click(timeout=8000)
        time.sleep(3)

        # 切换到“用密碼登入”
        try:
            page.locator(":text('用密碼登入')").first.click(timeout=6000)
            time.sleep(2)
        except Exception:
            log_fn("  未见“用密碼登入”入口，尝试直接找密码框")

        # 密码
        page.locator("input[type=password]").first.fill(password, timeout=10000)
        time.sleep(0.6)

        # 同意条款（如有）
        try:
            cbs = page.locator("input[type=checkbox]")
            for i in range(cbs.count()):
                cb = cbs.nth(i)
                if cb.is_visible() and not cb.is_checked():
                    cb.check(timeout=2000)
        except Exception:
            pass

        # 提交：弹窗里的绿色“登入”按钮（精确文本，避开右上角“登入 / 註冊”）
        submitted = False
        try:
            page.get_by_role("button", name="登入", exact=True).click(timeout=4000)
            submitted = True
        except Exception:
            try:
                page.locator(
                    "xpath=//*[normalize-space(text())='登入']"
                    "[not(ancestor::*[contains(.,'註冊')])]"
                ).last.click(timeout=4000)
                submitted = True
            except Exception as e:
                log_fn(f"  提交按钮点击失败: {e}")
        if not submitted:
            page.keyboard.press("Enter")
        time.sleep(6)

        body = page.evaluate("() => document.body.innerText")
        if "登入 / 註冊" in body:
            log_fn("登录后仍显示“登入 / 註冊”，可能失败（验证码/密码错误？）")
            return False
        log_fn("登录成功")
        return True
    except Exception as e:
        log_fn(f"登录异常: {e}")
        return False


def switch_currency_cny(page, log_fn) -> None:
    """把顶栏货币切换为 CNY。"""
    try:
        cur = page.evaluate(
            "() => { const m=document.body.innerText.match(/\\b(USD|CNY|HKD|AUD|JPY)\\b/); return m?m[1]:'?'; }"
        )
        if cur == "CNY":
            log_fn("货币已是 CNY")
            return
        page.locator("xpath=//*[normalize-space(text())='%s']" % cur).first.click(timeout=6000)
        time.sleep(1.5)
        page.locator(
            "xpath=(//*[normalize-space(text())='CNY' or contains(text(),'人民')])[1]"
        ).click(timeout=5000)
        time.sleep(3)
        cur2 = page.evaluate(
            "() => { const m=document.body.innerText.match(/\\b(USD|CNY|HKD|AUD|JPY)\\b/); return m?m[1]:'?'; }"
        )
        log_fn(f"货币切换结果: {cur2}")
    except Exception as e:
        log_fn(f"货币切换异常: {e}（继续，URL 已带 currency=CNY）")


def select_city(page, city: str, log_fn) -> None:
    """在列表页输入地区并选择联想首项。"""
    dest = page.locator(
        "xpath=//input[contains(@placeholder,'目的地') or contains(@placeholder,'城市') "
        "or contains(@placeholder,'地點') or contains(@placeholder,'酒店')]"
    ).first
    dest.wait_for(state="visible", timeout=20000)
    dest.click()
    dest.fill("")
    dest.type(city, delay=110)
    time.sleep(3)
    try:
        page.locator(
            f"xpath=(//li[contains(.,'{city}')] | //div[contains(@class,'item') and contains(.,'{city}')] "
            f"| //a[contains(.,'{city}')])[1]"
        ).click(timeout=6000)
        log_fn(f"已选择地区: {city}")
    except Exception:
        page.keyboard.press("Enter")
        log_fn("未命中联想项，按回车")
    time.sleep(5)


def fetch_detail_address(detail_page, url, log_fn) -> str:
    """打开酒店详情页，提取真实街道地址。失败返回空串。"""
    try:
        _goto_resilient(detail_page, url, log_fn)
        detail_page.wait_for_selector(".detail-info, .hotel-name, .map", timeout=20000)
        time.sleep(1.5)
        addr = detail_page.evaluate(r"""() => {
            const clean = s => (s||'').replace(/在地圖上顯示/g,'').replace(/\s+/g,' ').trim();
            let e = document.querySelector('.detail-info');
            if (e) { const t = clean(e.innerText); if (/省|市|區|区|路|街|號|号/.test(t)) return t; }
            // 兜底：找含“…省…市/…區…路號”的短文本
            for (const el of document.querySelectorAll('*')) {
                if (el.children.length <= 2) {
                    const t = clean(el.innerText);
                    if (t.length >= 6 && t.length < 80 && /(省|市).*(區|区|路|街|號|号)|,中國$|,中国$/.test(t))
                        return t;
                }
            }
            return '';
        }""")
        return (addr or "").strip()
    except Exception as e:
        log_fn(f"  取地址失败: {e}")
        return ""


def sort_by_price(page, log_fn) -> None:
    """选择“低價優先”排序（排序下拉：触发器 .sort-current，选项 li.sort-item）。"""
    try:
        # 打开排序下拉
        opened = False
        try:
            trig = page.locator(".sort-current").first
            if trig.count() > 0 and trig.is_visible():
                trig.click(timeout=5000)
                opened = True
        except Exception:
            pass
        if not opened:
            page.locator(
                "xpath=//*[normalize-space(text())='智能排序' or normalize-space(text())='低價優先']"
            ).first.click(timeout=5000)
        time.sleep(1.2)
        # 选择“低價優先”
        page.locator(
            "xpath=//li[contains(@class,'sort-item') and normalize-space(text())='低價優先']"
        ).first.click(timeout=5000)
        log_fn("已选择「低價優先」排序")
        time.sleep(3)
    except Exception as e:
        log_fn(f"排序异常: {e}")


# ============================================================
# 抓取核心
# ============================================================
def run_scraper(params: dict, log_callback=None) -> list:
    """
    params: city / in_date / out_date / adults / account / password /
            headless / max_scroll / threshold / coupon
    返回: [{"name","address","price","price_num","coupon","pay","url"}]
    """
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout

    results = []
    browser = None
    threshold = float(params.get("threshold", PAY_THRESHOLD_DEFAULT))   # 付款金额上限
    coupon = float(params.get("coupon", COUPON_DEFAULT))                 # 优惠券面值

    def _log(msg):
        (log_callback or print)(msg)

    try:
        with sync_playwright() as p:
            browser = launch_browser(p, params["headless"], _log)
            context = browser.new_context(
                viewport={"width": 1440, "height": 1000},
                locale="zh-Hant",
            )
            page = context.new_page()
            page.set_default_timeout(20000)

            # 0. 登录
            account = params.get("account", "").strip()
            password = params.get("password", "").strip()
            if account and password:
                do_login(page, account, password, _log)
            else:
                _log("未填写账号密码，以游客身份浏览（可能看不到优惠/折扣）")
                _goto_resilient(page, HOME_URL, _log)
                time.sleep(3)

            # 1. 切货币 CNY
            switch_currency_cny(page, _log)

            # 2. 进入列表页
            url = SEARCH_URL_TEMPLATE.format(
                in_date=params["in_date"],
                out_date=params["out_date"],
                adults=params["adults"],
            )
            _log(f"打开搜索页: {url}")
            _goto_resilient(page, url, _log)
            time.sleep(4)

            # 3. 输入地区
            select_city(page, params["city"], _log)

            # 4. 等待卡片
            try:
                page.wait_for_selector(".hotel-list-item", state="visible", timeout=25000)
            except PWTimeout:
                _log("未发现酒店卡片，结束")
                return results

            # 5. 低价排序
            sort_by_price(page, _log)
            try:
                page.wait_for_selector(".hotel-list-item", state="visible", timeout=15000)
            except PWTimeout:
                pass
            time.sleep(2)

            # 6. 滚动 + 点“查看更多住宿” 加载，逐批提取
            max_scroll = int(params.get("max_scroll", MAX_SCROLL_DEFAULT))
            seen = set()
            prev = -1
            stale = 0
            cutoff = threshold + coupon       # 房价 ≥ 此值 → 付款 ≥ 上限，不可能合格（已按价格升序）
            max_price_seen = 0.0
            for rnd in range(max_scroll):
                cards = page.evaluate(r"""() => {
                    return [...document.querySelectorAll('.hotel-list-item')].map(n => ({
                        name: ((n.querySelector('.hotel-name')||{}).textContent||'').trim(),
                        address: ((n.querySelector('.hotel-location-wrapper')||{}).innerText
                                  || (n.querySelector('.hotel-info')||{}).innerText || '').trim(),
                        price: ((n.querySelector('.sale-price')||{}).innerText||'').trim(),
                        tags: ((n.querySelector('.sale-tags')||{}).innerText||'').trim(),
                        href: n.getAttribute('href')
                              || ((n.querySelector('a')||{}).getAttribute
                                  ? (n.querySelector('a')||{}).getAttribute('href') : '') || ''
                    }));
                }""")

                for c in cards:
                    name = c.get("name") or ""
                    if not name:
                        continue
                    key = (name, c.get("price") or "")
                    if key in seen:
                        continue
                    seen.add(key)

                    # 付款金额 = 房价 − 优惠券；保留 付款金额 < 上限 的酒店
                    price_num = parse_price(c.get("price") or "")
                    if price_num is None:
                        continue
                    if price_num > max_price_seen:
                        max_price_seen = price_num
                    pay = max(0.0, price_num - coupon)
                    if pay >= threshold:
                        continue

                    href = (c.get("href") or "").strip()
                    full_url = ""
                    if href:
                        raw = href if href.startswith("http") else BASE_URL + href
                        # 去掉 traceToken 等冗长参数，保留干净的详情链接
                        m = re.search(r"hotelId=(\d+)", raw)
                        if m:
                            full_url = (
                                f"{BASE_URL}/zh-HK/hotel/hoteldetail?hotelId={m.group(1)}"
                                f"&inDate={params['in_date']}&outDate={params['out_date']}"
                                f"&adultsNumber={params['adults']}&currency=CNY"
                            )
                        else:
                            full_url = raw
                    addr = (c.get("address") or "").replace("\n", " ").strip() or "列表页未显示地址"
                    rec = {
                        "name": name,
                        "address": addr,
                        "price": (c.get("price") or "").strip(),   # 原始房价（如 連稅 CNY64）
                        "price_num": price_num,
                        "coupon": coupon,                          # 抵扣金额（优惠券面值）
                        "pay": pay,                                # 付款金额 = 房价 − 优惠券
                        "url": full_url,
                    }
                    results.append(rec)
                    _log(f"[命中] {name} | 抵扣 CNY{coupon:g} | 付款 CNY{pay:g}（房价 {rec['price']}）")

                total = len(seen)
                _log(f"--- 第 {rnd+1} 批：累计读取 {total} 家，符合 {len(results)} 家 ---")

                # 早停：已升序加载到房价 ≥ cutoff，更贵的不可能合格
                if max_price_seen >= cutoff:
                    _log(f"已加载到房价 ≥ CNY{cutoff:g}（更贵的不可能合格），停止加载")
                    break

                # 先滚动加载；若数量不再增长，点“查看更多住宿”
                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                time.sleep(2.5)
                if total == prev:
                    clicked = False
                    for sel in [".btn-loadmore", "text=查看更多住宿", "text=查看更多",
                                "text=載入更多", "text=加载更多"]:
                        try:
                            btn = page.locator(sel).first
                            if btn.count() > 0 and btn.is_visible():
                                btn.scroll_into_view_if_needed(timeout=3000)
                                btn.click(timeout=4000)
                                clicked = True
                                _log("点击“查看更多住宿”加载下一批")
                                time.sleep(3)
                                break
                        except Exception:
                            continue
                    if not clicked:
                        stale += 1
                        if stale >= 2:
                            _log("无更多酒店可加载，结束")
                            break
                else:
                    stale = 0
                prev = total

            # 7. 为符合条件的酒店补全真实街道地址（逐家进详情页）
            if params.get("fetch_address", True) and results:
                _log(f"开始补全 {len(results)} 家酒店的真实地址…")
                dp = context.new_page()
                dp.set_default_timeout(20000)
                for i, r in enumerate(results, 1):
                    if not r["url"]:
                        continue
                    addr = fetch_detail_address(dp, r["url"], _log)
                    if addr:
                        r["address"] = addr
                        _log(f"  ({i}/{len(results)}) {r['name']} -> {addr}")
                try:
                    dp.close()
                except Exception:
                    pass

            # 按付款金额排序，便于查看
            results.sort(key=lambda r: r["pay"])
            _log(f"抓取完成：共 {len(results)} 家付款金额 < CNY{threshold:g} 的酒店")
            return results
    finally:
        if browser is not None:
            try:
                browser.close()
            except Exception:
                pass


# ============================================================
# 文档输出
# ============================================================
def _header_lines(header: dict) -> list:
    return [
        f"生成时间：{header['generated_at']}",
        f"目标地区：{header['city']}",
        f"入住日期：{header['in_date']}",
        f"离店日期：{header['out_date']}",
        f"入住人数：{header['adults']}",
        f"优惠券面值：CNY {header['coupon']:g}",
        f"付款上限：CNY {header['threshold']:g}（仅保留 付款金额 < 此值 的酒店）",
        f"符合条件酒店数：{header['count']}",
    ]


def write_txt(results: list, path: str, header: dict) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write("HopeGoo 酒店筛选报告\n")
        f.write("=" * 60 + "\n")
        for line in _header_lines(header):
            f.write(line + "\n")
        f.write("=" * 60 + "\n\n")
        for idx, r in enumerate(results, 1):
            f.write(f"【{idx}】酒店名称：{to_simplified(r['name'])}\n")
            f.write(f"     酒店地址：{to_simplified(r['address'])}\n")
            f.write(f"     抵扣金额：CNY {r['coupon']:g}\n")
            f.write(f"     付款金额：CNY {r['pay']:g}（房价 {to_simplified(r['price'])}）\n")
            if r["url"]:
                f.write(f"     详情链接：{r['url']}\n")
            f.write("\n")
        if not results:
            f.write("（未找到符合条件的酒店）\n")


def write_docx(results: list, path: str, header: dict) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        txt_path = path.replace(".docx", ".txt")
        write_txt(results, txt_path, header)
        raise RuntimeError("未安装 python-docx，已回退保存为 TXT：" + txt_path)

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HopeGoo 酒店筛选报告")
    run.bold = True
    run.font.size = Pt(18)

    info = doc.add_paragraph()
    for line in _header_lines(header):
        info.add_run(line + "\n").font.size = Pt(11)

    doc.add_heading("酒店明细", level=2)
    for idx, r in enumerate(results, 1):
        p = doc.add_paragraph()
        p.add_run(f"【{idx}】{to_simplified(r['name'])}\n").bold = True
        p.add_run(f"    酒店地址：{to_simplified(r['address'])}\n")
        p.add_run(f"    抵扣金额：CNY {r['coupon']:g}\n")
        p.add_run(f"    付款金额：CNY {r['pay']:g}（房价 {to_simplified(r['price'])}）\n")
        if r["url"]:
            p.add_run(f"    详情链接：{r['url']}\n")

    if not results:
        doc.add_paragraph("（未找到符合条件的酒店）")
    doc.save(path)


def write_docx_table(results: list, path: str, header: dict) -> None:
    """表格格式的 Word 文档：每家酒店一行。"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        txt_path = path.replace(".docx", ".txt")
        write_txt(results, txt_path, header)
        raise RuntimeError("未安装 python-docx，已回退保存为 TXT：" + txt_path)

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HopeGoo 酒店筛选报告")
    run.bold = True
    run.font.size = Pt(18)

    info = doc.add_paragraph()
    for line in _header_lines(header):
        info.add_run(line + "\n").font.size = Pt(10)

    cols = ["序号", "酒店名称", "酒店地址", "抵扣金额(CNY)", "付款金额(CNY)", "房价"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        p = hdr[i].paragraphs[0]
        r = p.add_run(c)
        r.bold = True
        r.font.size = Pt(10)

    for idx, r in enumerate(results, 1):
        cells = table.add_row().cells
        vals = [str(idx), to_simplified(r["name"]), to_simplified(r["address"]), f"{r['coupon']:g}", f"{r['pay']:g}", to_simplified(r["price"])]
        for i, v in enumerate(vals):
            cp = cells[i].paragraphs[0]
            cp.add_run(v).font.size = Pt(9)

    if not results:
        doc.add_paragraph("（未找到符合条件的酒店）")
    doc.save(path)


def write_csv(results: list, path: str, header: dict) -> None:
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "序号", "酒店名称", "酒店地址", "抵扣金额(CNY)", "付款金额(CNY)", "房价", "详情链接",
            "生成时间", "目标地区", "入住日期", "离店日期", "入住人数", "优惠券面值", "付款上限",
        ])
        for idx, r in enumerate(results, 1):
            writer.writerow([
                idx, to_simplified(r["name"]), to_simplified(r["address"]), f"{r['coupon']:g}", f"{r['pay']:g}", to_simplified(r["price"]), r["url"],
                header["generated_at"], header["city"],
                header["in_date"], header["out_date"], header["adults"], header["coupon"], header["threshold"],
            ])


# ============================================================
# GUI
# ============================================================
class HopegooApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("HopeGoo 酒店筛选工具  |  按抵扣金额批量筛选")
        self.root.geometry("780x720")
        self.root.minsize(740, 660)
        self.accounts = load_accounts()
        self._build_ui()
        self._start_log_poller()

    def _build_ui(self) -> None:
        pad = {"padx": 10, "pady": 6}

        # 账号
        login_frm = ttk.LabelFrame(self.root, text="账号（必填，登录后才能看到优惠/折扣）")
        login_frm.pack(fill="x", **pad)

        # 第一行：已保存账号下拉切换
        ttk.Label(login_frm, text="已存账号：").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        self.saved_account_var = tk.StringVar(value="")
        self.account_combo = ttk.Combobox(login_frm, textvariable=self.saved_account_var,
                                           width=28, state="readonly")
        self.account_combo.grid(row=0, column=1, sticky="w", padx=4, pady=6)
        self.account_combo.bind("<<ComboboxSelected>>", self.on_pick_account)
        ttk.Button(login_frm, text="💾 保存当前", command=self.on_save_account).grid(row=0, column=2, padx=4, pady=6)
        ttk.Button(login_frm, text="🗑 删除", command=self.on_delete_account).grid(row=0, column=3, padx=4, pady=6, sticky="w")

        # 第二行：账号 / 密码输入
        ttk.Label(login_frm, text="账号 / 邮箱：").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        self.account_var = tk.StringVar(value="")
        ttk.Entry(login_frm, textvariable=self.account_var, width=30).grid(row=1, column=1, sticky="w", padx=4, pady=6)
        ttk.Label(login_frm, text="密码：").grid(row=1, column=2, sticky="e", padx=8, pady=6)
        self.password_var = tk.StringVar(value="")
        ttk.Entry(login_frm, textvariable=self.password_var, show="•", width=24).grid(row=1, column=3, sticky="w", padx=4, pady=6)

        self._refresh_account_combo()

        # 检索条件
        frm = ttk.LabelFrame(self.root, text="检索条件")
        frm.pack(fill="x", **pad)
        ttk.Label(frm, text="目标地区：").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        self.city_var = tk.StringVar(value="鶴崗")
        ttk.Entry(frm, textvariable=self.city_var, width=20).grid(row=0, column=1, sticky="w", padx=4, pady=6)

        ttk.Label(frm, text="入住日期 (YYYY-MM-DD)：").grid(row=0, column=2, sticky="e", padx=8, pady=6)
        today = datetime.now()
        self.in_date_var = tk.StringVar(value=(today + timedelta(days=7)).strftime("%Y-%m-%d"))
        ttk.Entry(frm, textvariable=self.in_date_var, width=16).grid(row=0, column=3, sticky="w", padx=4, pady=6)

        ttk.Label(frm, text="离店日期 (YYYY-MM-DD)：").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        self.out_date_var = tk.StringVar(value=(today + timedelta(days=8)).strftime("%Y-%m-%d"))
        ttk.Entry(frm, textvariable=self.out_date_var, width=20).grid(row=1, column=1, sticky="w", padx=4, pady=6)

        ttk.Label(frm, text="入住人数：").grid(row=1, column=2, sticky="e", padx=8, pady=6)
        self.adults_var = tk.StringVar(value="1")
        ttk.Spinbox(frm, from_=1, to=10, textvariable=self.adults_var, width=6).grid(row=1, column=3, sticky="w", padx=4, pady=6)

        # 运行选项
        adv = ttk.LabelFrame(self.root, text="运行选项")
        adv.pack(fill="x", **pad)

        ttk.Label(adv, text="优惠券面值 (CNY)：").grid(row=0, column=0, sticky="w", padx=10, pady=4)
        self.coupon_var = tk.StringVar(value=str(COUPON_DEFAULT))
        ttk.Spinbox(adv, from_=0, to=9999, textvariable=self.coupon_var, width=6).grid(row=0, column=1, sticky="w", padx=4, pady=4)

        ttk.Label(adv, text="付款上限 (CNY)：").grid(row=0, column=2, sticky="e", padx=10, pady=4)
        self.threshold_var = tk.StringVar(value=str(PAY_THRESHOLD_DEFAULT))
        ttk.Spinbox(adv, from_=1, to=9999, textvariable=self.threshold_var, width=6).grid(row=0, column=3, sticky="w", padx=4, pady=4)

        ttk.Label(adv, text="最多加载批次：").grid(row=1, column=0, sticky="w", padx=10, pady=4)
        self.max_scroll_var = tk.StringVar(value=str(MAX_SCROLL_DEFAULT))
        ttk.Spinbox(adv, from_=1, to=500, textvariable=self.max_scroll_var, width=5).grid(row=1, column=1, sticky="w", padx=4, pady=4)

        self.headless_var = tk.BooleanVar(value=HEADLESS_DEFAULT)
        ttk.Checkbutton(adv, text="无头模式（不显示浏览器窗口）", variable=self.headless_var).grid(
            row=1, column=2, columnspan=2, sticky="w", padx=10, pady=4)

        self.fetch_addr_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(adv, text="补全真实地址（逐家进详情页，较慢）", variable=self.fetch_addr_var).grid(
            row=4, column=0, columnspan=2, sticky="w", padx=10, pady=4)

        ttk.Label(adv, text="输出文档格式：").grid(row=2, column=0, sticky="w", padx=10, pady=4)
        self.format_var = tk.StringVar(value="表格(Word)")
        ttk.Combobox(adv, textvariable=self.format_var,
                     values=["表格(Word)", "DOCX", "TXT", "CSV"], width=12, state="readonly").grid(
            row=2, column=1, sticky="w", padx=4, pady=4)

        self.output_dir_var = tk.StringVar(
            value=os.path.join(os.path.dirname(os.path.abspath(__file__)), "output"))
        ttk.Label(adv, text="输出目录：").grid(row=2, column=2, sticky="e", padx=10, pady=4)
        ttk.Entry(adv, textvariable=self.output_dir_var, width=34).grid(row=2, column=3, sticky="we", padx=4, pady=4)
        ttk.Label(adv, text="（文件自动按「地区_日期时间」命名，每次生成新文档）",
                  foreground="#888").grid(row=3, column=2, columnspan=2, sticky="e", padx=10)
        adv.columnconfigure(3, weight=1)

        # 按钮栏
        btn_bar = ttk.Frame(self.root)
        btn_bar.pack(fill="x", **pad)
        self.start_btn = ttk.Button(btn_bar, text="▶ 开始筛选", command=self.on_start)
        self.start_btn.pack(side="left", padx=6)
        ttk.Button(btn_bar, text="📂 选择输出目录", command=self.on_choose_path).pack(side="left", padx=6)
        self.open_btn = ttk.Button(btn_bar, text="📄 打开输出文件", command=self.on_open_output, state="disabled")
        self.open_btn.pack(side="left", padx=6)
        ttk.Button(btn_bar, text="🧹 清空日志", command=self.on_clear_log).pack(side="left", padx=6)
        self.last_output_path = None

        # 日志
        log_frm = ttk.LabelFrame(self.root, text="运行日志")
        log_frm.pack(fill="both", expand=True, **pad)
        self.log_text = scrolledtext.ScrolledText(log_frm, wrap="word", height=18, font=("Menlo", 11))
        self.log_text.pack(fill="both", expand=True, padx=6, pady=6)

        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(self.root, textvariable=self.status_var, anchor="w", relief="sunken").pack(fill="x")

    # ---------- 事件 ----------
    def on_choose_path(self) -> None:
        d = filedialog.askdirectory(title="选择输出目录", initialdir=self.output_dir_var.get() or ".")
        if d:
            self.output_dir_var.set(d)

    def on_clear_log(self) -> None:
        self.log_text.delete("1.0", "end")

    def on_open_output(self) -> None:
        try:
            open_file(self.last_output_path)
        except Exception as e:
            messagebox.showerror("打开失败", f"无法打开文件：{e}")

    # ---------- 账号管理 ----------
    def _refresh_account_combo(self, select: str = None) -> None:
        names = [a["account"] for a in self.accounts]
        self.account_combo["values"] = names
        if select and select in names:
            self.saved_account_var.set(select)
        elif names and not self.saved_account_var.get():
            # 启动时默认带出第一个账号
            self.saved_account_var.set(names[0])
            self._fill_from_account(names[0])

    def _fill_from_account(self, account: str) -> None:
        for a in self.accounts:
            if a["account"] == account:
                self.account_var.set(a["account"])
                self.password_var.set(a.get("password", ""))
                return

    def on_pick_account(self, event=None) -> None:
        self._fill_from_account(self.saved_account_var.get())

    def on_save_account(self) -> None:
        acc = self.account_var.get().strip()
        pwd = self.password_var.get().strip()
        if not acc or not pwd:
            messagebox.showerror("无法保存", "请先填写账号和密码再保存")
            return
        for a in self.accounts:                       # 已存在则更新密码
            if a["account"] == acc:
                a["password"] = pwd
                break
        else:
            self.accounts.append({"account": acc, "password": pwd})
        try:
            save_accounts(self.accounts)
        except Exception as e:
            messagebox.showerror("保存失败", str(e))
            return
        self._refresh_account_combo(select=acc)
        messagebox.showinfo("已保存", f"账号已保存：{acc}\n（明文存于 accounts.json）")

    def on_delete_account(self) -> None:
        acc = self.saved_account_var.get().strip()
        if not acc:
            return
        if not messagebox.askyesno("删除账号", f"确定删除已存账号：{acc}？"):
            return
        self.accounts = [a for a in self.accounts if a["account"] != acc]
        try:
            save_accounts(self.accounts)
        except Exception as e:
            messagebox.showerror("删除失败", str(e))
            return
        self.saved_account_var.set("")
        self._refresh_account_combo()

    def on_start(self) -> None:
        city = self.city_var.get().strip()
        in_date = self.in_date_var.get().strip()
        out_date = self.out_date_var.get().strip()
        adults = self.adults_var.get().strip()
        if not city:
            messagebox.showerror("参数错误", "请填写目标地区")
            return
        try:
            datetime.strptime(in_date, "%Y-%m-%d")
            datetime.strptime(out_date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("参数错误", "日期格式应为 YYYY-MM-DD")
            return
        if in_date >= out_date:
            messagebox.showerror("参数错误", "离店日期必须晚于入住日期")
            return
        try:
            int(adults)
            threshold = float(self.threshold_var.get())
            coupon = float(self.coupon_var.get())
        except ValueError:
            messagebox.showerror("参数错误", "入住人数 / 优惠券面值 / 付款上限必须为数字")
            return

        params = {
            "city": city, "in_date": in_date, "out_date": out_date, "adults": int(adults),
            "account": self.account_var.get().strip(), "password": self.password_var.get().strip(),
            "headless": self.headless_var.get(),
            "max_scroll": int(self.max_scroll_var.get() or MAX_SCROLL_DEFAULT),
            "threshold": threshold,
            "coupon": coupon,
            "fetch_address": self.fetch_addr_var.get(),
            "output_dir": self.output_dir_var.get().strip()
                          or os.path.join(os.path.dirname(os.path.abspath(__file__)), "output"),
            "output_format": self.format_var.get(),
        }

        if not params["account"] or not params["password"]:
            if not messagebox.askyesno("未填写账号", "未填写账号密码，将以游客身份浏览，可能看不到优惠。仍要继续吗？"):
                return

        self.start_btn.config(state="disabled")
        self.status_var.set("正在抓取，请稍候…")
        log("=" * 50)
        log(f"开始筛选：地区={city} / {in_date}~{out_date} / {adults} 人 / 优惠券CNY{coupon:g} / 付款<CNY{threshold:g}")
        threading.Thread(target=self._worker, args=(params,), daemon=True).start()

    def _worker(self, params: dict) -> None:
        try:
            results = run_scraper(params, log_callback=log)
            header = {
                "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "city": params["city"], "in_date": params["in_date"],
                "out_date": params["out_date"], "adults": params["adults"],
                "threshold": float(params["threshold"]), "coupon": float(params["coupon"]),
                "count": len(results),
            }
            fmt = params["output_format"]
            ext = {"表格(Word)": "docx", "DOCX": "docx", "TXT": "txt", "CSV": "csv"}.get(fmt, "docx")
            out_dir = params["output_dir"]
            os.makedirs(out_dir, exist_ok=True)
            # 文件名：地区_日期时间，每次都是新文档
            safe_city = re.sub(r'[\\/:*?"<>|\s]+', "", params["city"]) or "酒店"
            stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
            path = os.path.join(out_dir, f"{safe_city}_{stamp}.{ext}")

            if fmt == "表格(Word)":
                try:
                    write_docx_table(results, path, header)
                except RuntimeError as e:
                    log(str(e)); path = path.replace(".docx", ".txt")
            elif fmt == "DOCX":
                try:
                    write_docx(results, path, header)
                except RuntimeError as e:
                    log(str(e)); path = path.replace(".docx", ".txt")
            elif fmt == "TXT":
                write_txt(results, path, header)
            else:
                write_csv(results, path, header)

            log(f"文档已生成：{path}")
            self.last_output_path = path
            self.root.after(0, lambda: self.open_btn.config(state="normal"))
            # 生成后自动打开
            try:
                open_file(path)
            except Exception as e:
                log(f"自动打开失败（可手动点“打开输出文件”）：{e}")
            self.root.after(0, lambda: messagebox.showinfo(
                "完成", f"筛选完成！\n共 {len(results)} 家符合条件的酒店。\n文档：{path}"))
            self.root.after(0, lambda: self.status_var.set(f"完成，共 {len(results)} 家"))
        except Exception as e:
            log(f"执行出错：{e}")
            self.root.after(0, lambda: messagebox.showerror("错误", f"执行出错：{e}"))
            self.root.after(0, lambda: self.status_var.set("运行失败"))
        finally:
            self.root.after(0, lambda: self.start_btn.config(state="normal"))

    def _start_log_poller(self) -> None:
        def poller():
            while True:
                try:
                    msg = LOG_QUEUE.get_nowait()
                except queue.Empty:
                    break
                self.log_text.insert("end", msg + "\n")
                self.log_text.see("end")
            self.root.after(300, poller)
        self.root.after(300, poller)


def main() -> None:
    root = tk.Tk()
    try:
        style = ttk.Style(root)
        if sys.platform == "darwin":
            style.theme_use("aqua")
        elif sys.platform.startswith("win"):
            style.theme_use("vista")
    except Exception:
        pass
    HopegooApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
